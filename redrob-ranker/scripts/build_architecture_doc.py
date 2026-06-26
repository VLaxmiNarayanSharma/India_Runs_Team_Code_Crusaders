#!/usr/bin/env python3
"""Generate Architecture_working.docx for the Redrob ranker project."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_PATH = Path(__file__).resolve().parents[1] / "Architecture_working.docx"


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    doc.add_paragraph(text, style="List Bullet")


def build_document() -> Document:
    doc = Document()

    title = doc.add_heading("Redrob Ranker — Project Architecture", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This document describes the architecture of the AI candidate ranking pipeline built "
        "for the Redrob INDIA RUNS hackathon (Intelligent Candidate Discovery & Ranking Challenge)."
    )

    # 1. High-level idea
    doc.add_heading("1. High-Level Idea", level=1)
    doc.add_paragraph(
        "This project is an AI candidate ranking pipeline. It reads a job description and "
        "100,000 candidate profiles, then outputs a top-100 shortlist in submission.csv — "
        "the way a strong recruiter would, not via keyword matching."
    )
    doc.add_paragraph(
        "Traditional hiring tools fail because they rank on keyword overlap (e.g. Python, ML, RAG). "
        "The Redrob job description explicitly warns that this is a honeypot trap in the dataset."
    )
    doc.add_paragraph("The system uses a multi-stage funnel:")
    add_code_block(
        doc,
        "100,000 candidates\n"
        "    → understand the JD\n"
        "    → build rich candidate representations\n"
        "    → fast hybrid retrieval (→ 3,000)\n"
        "    → interpretable scoring + honeypot filtering (→ 100)\n"
        "    → template reasoning strings\n"
        "    → submission.csv",
    )

    # 2. Pipeline stages
    doc.add_heading("2. End-to-End Pipeline (6 Stages)", level=1)

    stages = [
        (
            "Stage 1 — JD Understanding (src/jd_parser.py)",
            "Parse the job description into structured requirements — not just keywords.",
            "Input: job_description.docx",
            "Output: data/precomputed/jd_requirements.json",
            [
                "Role title (e.g. Senior AI Engineer - Founding Team)",
                "Role family (search_retrieval_ranking)",
                "Experience band (5–9 years)",
                "Must-haves (embeddings, hybrid retrieval, evaluation, Python)",
                "Locations (Pune, Noida, Hyderabad, Mumbai, Delhi)",
                "Disqualifiers (consulting-only, keyword traps, pure research)",
                "Canonical JD text for retrieval/embedding",
            ],
            "Parsing runs once offline; ranking reuses the JSON artifact with no network calls.",
        ),
        (
            "Stage 2 — Candidate Representation (src/candidate_features.py)",
            "Turn each raw profile into searchable text plus structured signals.",
            "Input: candidates.jsonl + jd_requirements.json",
            "Output: candidate_features_full.jsonl (~306 MB for 100K candidates)",
            [
                "Canonical text: headline + summary + title + career descriptions + top skills",
                "role_fit: title similarity, % career in ML/IR",
                "skills: depth, verification gap, keyword stuffing ratio",
                "career_coherence: title alignment, progression, description uniqueness",
                "authenticity: copied bios, title–skill mismatch",
                "hireability: experience fit, location, response rate, recency",
                "behavioral: GitHub, recruiter saves, interview completion",
            ],
            None,
        ),
        (
            "Stage 3 — Hybrid Retrieval (src/retrieval.py)",
            "Narrow 100,000 candidates to ~3,000 quickly on CPU.",
            "Input: candidate_features_full.jsonl + jd_requirements.json",
            "Output: retrieved_top3000.jsonl, candidate_embeddings.npy",
            [
                "Dense: all-MiniLM-L6-v2 embeddings (cached)",
                "Lexical: BM25 (rank_bm25)",
                "Formula: 0.6 × cosine_sim + 0.4 × BM25 (min-max normalized)",
                "Re-run time: ~57 seconds with cached embeddings",
            ],
            "Dense search finds semantic fit; BM25 catches exact terms (NDCG, BM25, Milvus).",
        ),
        (
            "Stages 4–5 — Scoring + Honeypot Filtering (src/scoring.py, src/honeypot.py)",
            "Rank the 3,000 retrieved candidates with interpretable, recruiter-like logic.",
            "Input: retrieved_top3000.jsonl + features + raw profiles",
            "Output: scored_top100.jsonl, honeypot_assessments.jsonl, submission.csv",
            [
                "semantic_fit (0.12): retrieval + title + skill coverage",
                "career_trajectory (0.38): dominant signal — beats keyword filters",
                "skill_depth_trust (0.12): depth over count",
                "behavioral_hireability (0.18): response rate, recency, notice period",
                "inconsistency_penalty (-0.10): copied bios, mismatches",
                "keyword_stuffing_penalty (-0.15): honeypot trap",
                "honeypot_penalty (-0.20): explicit anti-trap rules",
            ],
            "2,585 honeypots detected in the 3K pool; top 100 are real AI/search engineers.",
        ),
        (
            "Stage 6 — Reasoning Generation (src/reasoning.py)",
            "Produce recruiter-trustworthy explanations with no LLM at runtime.",
            "Input: scored_top100.jsonl + raw profiles",
            "Output: submission.csv (refreshed reasoning column)",
            [
                'Example: "Senior AI Engineer with 7.8 yrs; hybrid retrieval + LTR in production; '
                '11 verified core skills; response rate 0.76."',
                "Technical highlights detected via pattern matching on career text and skills",
            ],
            None,
        ),
    ]

    for title, goal, inp, out, bullets, note in stages:
        doc.add_heading(title, level=2)
        doc.add_paragraph(goal)
        doc.add_paragraph(inp)
        doc.add_paragraph(out)
        for b in bullets:
            add_bullet(doc, b)
        if note:
            p = doc.add_paragraph()
            p.add_run("Note: ").bold = True
            p.add_run(note)

    doc.add_paragraph()
    doc.add_heading("Scoring Formula", level=2)
    add_code_block(
        doc,
        "total_score = 0.12 × semantic_fit\n"
        "            + 0.38 × career_trajectory\n"
        "            + 0.12 × skill_depth_trust\n"
        "            + 0.18 × behavioral_hireability\n"
        "            - 0.10 × inconsistency_penalty\n"
        "            - 0.15 × keyword_stuffing_penalty\n"
        "            - 0.20 × honeypot_penalty",
    )

    doc.add_heading("Honeypot Rules (Stage 5)", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Rule"
    hdr[1].text = "Penalty"
    rules = [
        ("Many AI skills + non-technical career", "Large (0.85)"),
        ("Copied job descriptions (Jaccard > 0.7)", "Medium (0.55)"),
        ("Title ≠ dominant career theme", "Medium (0.55)"),
        ("High endorsements + low assessments", "skill_trust × 0.5"),
    ]
    for rule, penalty in rules:
        row = table.add_row().cells
        row[0].text = rule
        row[1].text = penalty

    # 3. File structure
    doc.add_heading("3. Project File Structure", level=1)
    add_code_block(
        doc,
        "redrob-ranker/\n"
        "├── rank.py                    # Main entry — runs full pipeline\n"
        "├── submission.csv             # Final top-100 output\n"
        "├── requirements.txt\n"
        "├── scripts/\n"
        "│   ├── run_stage1.py          # JD parsing\n"
        "│   ├── run_stage2.py          # Feature extraction\n"
        "│   ├── run_stage3.py          # Hybrid retrieval\n"
        "│   ├── run_stage4.py          # Scoring only (optional)\n"
        "│   ├── run_stage5.py          # Scoring + honeypots + CSV\n"
        "│   └── run_stage6.py          # Reasoning refresh\n"
        "├── src/\n"
        "│   ├── jd_parser.py           # Stage 1\n"
        "│   ├── candidate_features.py  # Stage 2\n"
        "│   ├── retrieval.py           # Stage 3\n"
        "│   ├── scoring.py             # Stage 4\n"
        "│   ├── honeypot.py            # Stage 5\n"
        "│   ├── reasoning.py           # Stage 6\n"
        "│   ├── candidate_loader.py\n"
        "│   ├── constants.py\n"
        "│   └── paths.py\n"
        "└── data/precomputed/\n"
        "    ├── jd_requirements.json\n"
        "    ├── candidate_features_full.jsonl\n"
        "    ├── candidate_embeddings.npy\n"
        "    ├── retrieved_top3000.jsonl\n"
        "    ├── honeypot_assessments.jsonl\n"
        "    └── scored_top100.jsonl",
    )

    # 4. Design principles
    doc.add_heading("4. Design Principles", level=1)
    principles = [
        (
            "Funnel architecture",
            "Scoring 100K with heavy logic would be slow. Retrieval first, then deep scoring on 3K keeps runtime practical.",
        ),
        (
            "Career history > skill count",
            "career_trajectory weight (0.38) is the largest positive signal. This beats HR Manager with 9 AI core skills honeypots.",
        ),
        (
            "Offline precomputation",
            "Embeddings and features are cached. Hackathon requires no network during ranking — this design complies.",
        ),
        (
            "Interpretable over black-box",
            "Every score has named components. Reasoning strings explain why a candidate ranked high.",
        ),
        (
            "Explicit honeypot handling",
            "The JD says keyword matching is wrong. Stage 5 encodes that directly.",
        ),
    ]
    for name, desc in principles:
        p = doc.add_paragraph()
        p.add_run(f"{name}. ").bold = True
        p.add_run(desc)

    # 5. Data flow
    doc.add_heading("5. Data Flow Summary", level=1)
    add_code_block(
        doc,
        "job_description.docx\n"
        "        │\n"
        "        ▼\n"
        "  jd_requirements.json ─────────────────────────────┐\n"
        "        │                                          │\n"
        "candidates.jsonl                                   │\n"
        "        │                                          │\n"
        "        ▼                                          ▼\n"
        "candidate_features_full.jsonl              retrieval.py\n"
        "        │                                    (embed + BM25)\n"
        "        │                                          │\n"
        "        └──────────────┬───────────────────────────┘\n"
        "                       ▼\n"
        "              retrieved_top3000.jsonl\n"
        "                       │\n"
        "                       ▼\n"
        "         scoring.py + honeypot.py\n"
        "                       │\n"
        "                       ▼\n"
        "              scored_top100.jsonl\n"
        "                       │\n"
        "                       ▼\n"
        "              reasoning.py\n"
        "                       │\n"
        "                       ▼\n"
        "               submission.csv",
    )

    # 6. How to run
    doc.add_heading("6. How to Run", level=1)
    doc.add_paragraph("Full pipeline (skip stages 1–3 if artifacts already exist):")
    add_code_block(doc, "python redrob-ranker/rank.py --skip-stage1 --skip-stage2 --skip-stage3")
    doc.add_paragraph("Individual stages:")
    add_code_block(
        doc,
        "python redrob-ranker/scripts/run_stage1.py   # Parse JD\n"
        "python redrob-ranker/scripts/run_stage2.py   # Extract features\n"
        "python redrob-ranker/scripts/run_stage3.py   # Retrieve top 3K\n"
        "python redrob-ranker/scripts/run_stage5.py   # Score + honeypots\n"
        "python redrob-ranker/scripts/run_stage6.py   # Refresh reasoning",
    )
    doc.add_paragraph("Validate submission:")
    add_code_block(doc, "python validate_submission.py redrob-ranker/submission.csv")

    # 7. Good vs bad
    doc.add_heading("7. Good vs Bad Candidates (In This System)", level=1)
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Good fit (ranks high)"
    t2.rows[0].cells[1].text = "Bad fit (penalized)"
    comparisons = [
        ("Senior AI / Search / NLP Engineer", "HR Manager with many AI skills"),
        ("Career shows retrieval, ranking, production ML", "Copied descriptions across jobs"),
        ("Verified skills + assessments", "High endorsements, low assessments"),
        ("Active, responsive on platform", "Inactive, low response rate"),
        ("Hybrid retrieval + LTR in career text", "Keyword-stuffed skill list only"),
    ]
    for good, bad in comparisons:
        row = t2.add_row().cells
        row[0].text = good
        row[1].text = bad

    doc.add_paragraph()
    doc.add_paragraph(
        "Document generated for the Redrob INDIA RUNS Data & AI Challenge.",
        style="Intense Quote",
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
